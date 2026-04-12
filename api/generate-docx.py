from http.server import BaseHTTPRequestHandler
import json
import os
import io
import sys
from docx import Document

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), '..', 'template.docx')

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        para = cell.paragraphs[0]
        if para.runs:
            para.runs[0].text = text
        else:
            para.add_run(text)
    else:
        cell.add_paragraph(text)

def replace_highlighted_text(cell, new_text):
    from docx.oxml.ns import qn
    for para in cell.paragraphs:
        highlighted_runs = []
        for run in para.runs:
            rpr = run._r.find(qn('w:rPr'))
            if rpr is not None:
                hl = rpr.find(qn('w:highlight'))
                if hl is not None and hl.get(qn('w:val')) == 'yellow':
                    highlighted_runs.append(run)
        if highlighted_runs:
            highlighted_runs[0].text = new_text
            for r in highlighted_runs[1:]:
                r.text = ""
            break

def fill_doctors(table, doctors):
    dr_rows = [(18, 19), (21, 22), (24, 25), (27, 28), (30, 31)]
    for i, (name_row, dept_row) in enumerate(dr_rows):
        if i < len(doctors):
            dr = doctors[i]
            set_cell_text(table.rows[name_row].cells[4], dr.get('name', ''))
            set_cell_text(table.rows[dept_row].cells[3], dr.get('dept', ''))

def build_complaint_text(data):
    return (
        f"(1) 対応窓口\n"
        f"窓口: {data.get('cperson', '')}\n"
        f"受付住所: {data.get('caddr2', '')}\n"
        f"電話番号: {data.get('cphone', '')}\n"
        f"受付時間：休診日を除く {data.get('chours', '')}\n"
        f"(2) 対応手順\n"
        f"苦情や問合せを受けた場合、対応窓口の担当者はその詳細を管理者に伝え、管理者は対応を検討した上で適切に処理する。"
        f"営業時間外に備えて患者へ緊急連絡先を別途案内する。"
    )

def fill_template(data):
    doc = Document(TEMPLATE_PATH)
    t = doc.tables

    # 医療機関情報
    set_cell_text(t[0].rows[0].cells[2], data.get('cname', ''))
    set_cell_text(t[0].rows[1].cells[2], data.get('caddr', ''))
    set_cell_text(t[0].rows[2].cells[2], data.get('dir', ''))

    # 治療名称・対象疾患
    replace_highlighted_text(t[1].rows[0].cells[1], data.get('tname', ''))
    replace_highlighted_text(t[1].rows[4].cells[1], data.get('targetDisease', ''))
    replace_highlighted_text(t[1].rows[5].cells[1], data.get('targetDisease', ''))

    # 人員
    set_cell_text(t[2].rows[1].cells[4],  data.get('respName', ''))
    set_cell_text(t[2].rows[2].cells[4],  data.get('respOrg', ''))
    set_cell_text(t[2].rows[3].cells[4],  data.get('respDept', ''))
    set_cell_text(t[2].rows[4].cells[4],  data.get('respZip', ''))
    set_cell_text(t[2].rows[5].cells[4],  data.get('respAddr', ''))
    set_cell_text(t[2].rows[6].cells[4],  data.get('respPhone', ''))
    set_cell_text(t[2].rows[7].cells[4],  data.get('respEmail', ''))
    set_cell_text(t[2].rows[8].cells[4],  data.get('respRole', ''))
    set_cell_text(t[2].rows[9].cells[4],  data.get('staffName', ''))
    set_cell_text(t[2].rows[10].cells[4], data.get('staffOrg', ''))
    set_cell_text(t[2].rows[11].cells[4], data.get('staffDept', ''))
    set_cell_text(t[2].rows[12].cells[4], data.get('staffZip', ''))
    set_cell_text(t[2].rows[13].cells[4], data.get('staffAddr', ''))
    set_cell_text(t[2].rows[14].cells[4], data.get('staffPhone', ''))
    set_cell_text(t[2].rows[15].cells[4], data.get('staffFax', ''))
    set_cell_text(t[2].rows[16].cells[4], data.get('staffEmail', ''))
    fill_doctors(t[2], data.get('doctors', []))
    set_cell_text(t[2].rows[33].cells[3], data.get('emergencyText', ''))

    # 採取方法
    set_cell_text(t[5].rows[1].cells[1], data.get('collectionText', ''))

    # 投与方法・製造施設
    replace_highlighted_text(t[6].rows[2].cells[2], data.get('adminText', ''))
    set_cell_text(t[6].rows[4].cells[2], data.get('mfrName', ''))
    set_cell_text(t[6].rows[5].cells[2], data.get('fname', ''))
    set_cell_text(t[6].rows[6].cells[3], 'FC' + data.get('fcnum', ''))

    # 補償内容
    ins = data.get('insuranceText', '')
    if ins:
        set_cell_text(t[12].rows[2].cells[2], ins)
        set_cell_text(t[12].rows[5].cells[2], ins)

    # 委員会
    comm = data.get('comm', 'JSCSF再生医療等委員会　NA8230002')
    parts = comm.split('　')
    set_cell_text(t[13].rows[0].cells[1], parts[0] if parts else comm)
    set_cell_text(t[13].rows[1].cells[1], parts[1] if len(parts) > 1 else '')

    # 苦情窓口
    set_cell_text(t[14].rows[3].cells[1], build_complaint_text(data))

    # メモリ上に保存して返す
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        content_length = int(self.headers.get('Content-Length', 0))
        body = self.rfile.read(content_length)

        # CORS
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        self.send_header('Content-Disposition', 'attachment; filename="様式1-2.docx"')

        try:
            data = json.loads(body)
            docx_bytes = fill_template(data)
            self.send_header('Content-Length', str(len(docx_bytes)))
            self.end_headers()
            self.wfile.write(docx_bytes)
        except Exception as e:
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
